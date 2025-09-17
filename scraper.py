import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
from PIL import Image
import pytesseract
import os
import json
import re
from urllib.parse import urljoin, urlparse
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException

class DataScraper:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
    def scrape_website(self, url):
        """Scrape general website content"""
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract key information
            data = {
                'url': url,
                'title': soup.find('title').get_text().strip() if soup.find('title') else '',
                'description': '',
                'content': '',
                'links': [],
                'images': [],
                'social_links': {},
                'contact_info': {},
                'scraped_at': time.time()
            }
            
            # Meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc:
                data['description'] = meta_desc.get('content', '')
            
            # Main content (try to get the most relevant content)
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            if main_content:
                data['content'] = main_content.get_text().strip()[:2000]  # Limit content
            
            # Extract links
            for link in soup.find_all('a', href=True):
                href = link.get('href')
                if href:
                    full_url = urljoin(url, href)
                    data['links'].append({
                        'url': full_url,
                        'text': link.get_text().strip()
                    })
            
            # Extract images
            for img in soup.find_all('img', src=True):
                src = img.get('src')
                if src:
                    full_url = urljoin(url, src)
                    data['images'].append({
                        'url': full_url,
                        'alt': img.get('alt', '')
                    })
            
            # Extract social media links
            social_patterns = {
                'linkedin': r'linkedin\.com',
                'twitter': r'twitter\.com|x\.com',
                'facebook': r'facebook\.com',
                'instagram': r'instagram\.com',
                'youtube': r'youtube\.com',
                'tiktok': r'tiktok\.com'
            }
            
            for platform, pattern in social_patterns.items():
                for link in data['links']:
                    if re.search(pattern, link['url'], re.IGNORECASE):
                        data['social_links'][platform] = link['url']
                        break
            
            # Extract contact information
            text_content = soup.get_text()
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            phone_pattern = r'(\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
            
            emails = re.findall(email_pattern, text_content)
            phones = re.findall(phone_pattern, text_content)
            
            data['contact_info'] = {
                'emails': list(set(emails)),
                'phones': list(set(phones))
            }
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'url': url}
    
    def scrape_social_media(self, platform, username_or_url):
        """Scrape social media profiles"""
        try:
            if platform == 'linkedin':
                return self._scrape_linkedin(username_or_url)
            elif platform == 'twitter':
                return self._scrape_twitter(username_or_url)
            elif platform == 'instagram':
                return self._scrape_instagram(username_or_url)
            elif platform == 'facebook':
                return self._scrape_facebook(username_or_url)
            elif platform == 'youtube':
                return self._scrape_youtube(username_or_url)
            elif platform == 'tiktok':
                return self._scrape_tiktok(username_or_url)
            else:
                return {'error': f'Unsupported platform: {platform}'}
        except Exception as e:
            return {'error': str(e), 'platform': platform, 'username': username_or_url}
    
    def _scrape_linkedin(self, username_or_url):
        """Scrape LinkedIn profile/company page"""
        # LinkedIn requires special handling due to anti-bot measures
        # This is a simplified version - in production, you'd need proper authentication
        try:
            if not username_or_url.startswith('http'):
                url = f"https://linkedin.com/in/{username_or_url}"
            else:
                url = username_or_url
                
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'platform': 'linkedin',
                'url': url,
                'name': '',
                'headline': '',
                'about': '',
                'followers': '',
                'scraped_at': time.time()
            }
            
            # Extract basic info (LinkedIn's structure changes frequently)
            name_elem = soup.find('h1', class_='text-heading-xlarge')
            if name_elem:
                data['name'] = name_elem.get_text().strip()
            
            headline_elem = soup.find('div', class_='text-body-medium')
            if headline_elem:
                data['headline'] = headline_elem.get_text().strip()
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'platform': 'linkedin'}
    
    def _scrape_twitter(self, username_or_url):
        """Scrape Twitter profile"""
        try:
            if not username_or_url.startswith('http'):
                url = f"https://twitter.com/{username_or_url}"
            else:
                url = username_or_url
                
            # Use Selenium for Twitter due to dynamic content
            options = Options()
            options.add_argument('--headless')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            
            driver = webdriver.Chrome(options=options)
            driver.get(url)
            
            data = {
                'platform': 'twitter',
                'url': url,
                'name': '',
                'username': '',
                'bio': '',
                'followers': '',
                'following': '',
                'tweets': '',
                'scraped_at': time.time()
            }
            
            try:
                # Wait for page to load
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, '[data-testid="UserName"]'))
                )
                
                # Extract profile info
                name_elem = driver.find_element(By.CSS_SELECTOR, '[data-testid="UserName"]')
                data['name'] = name_elem.text if name_elem else ''
                
                username_elem = driver.find_element(By.CSS_SELECTOR, '[data-testid="UserName"] + div')
                data['username'] = username_elem.text if username_elem else ''
                
                bio_elem = driver.find_element(By.CSS_SELECTOR, '[data-testid="UserDescription"]')
                data['bio'] = bio_elem.text if bio_elem else ''
                
            except TimeoutException:
                data['error'] = 'Timeout loading Twitter profile'
            finally:
                driver.quit()
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'platform': 'twitter'}
    
    def _scrape_instagram(self, username_or_url):
        """Scrape Instagram profile"""
        try:
            if not username_or_url.startswith('http'):
                url = f"https://instagram.com/{username_or_url}"
            else:
                url = username_or_url
                
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'platform': 'instagram',
                'url': url,
                'name': '',
                'username': '',
                'bio': '',
                'followers': '',
                'following': '',
                'posts': '',
                'scraped_at': time.time()
            }
            
            # Instagram uses JSON-LD structured data
            json_scripts = soup.find_all('script', type='application/ld+json')
            for script in json_scripts:
                try:
                    json_data = json.loads(script.string)
                    if '@type' in json_data and json_data['@type'] == 'Person':
                        data['name'] = json_data.get('name', '')
                        data['bio'] = json_data.get('description', '')
                        break
                except:
                    continue
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'platform': 'instagram'}
    
    def _scrape_facebook(self, username_or_url):
        """Scrape Facebook page"""
        try:
            if not username_or_url.startswith('http'):
                url = f"https://facebook.com/{username_or_url}"
            else:
                url = username_or_url
                
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'platform': 'facebook',
                'url': url,
                'name': '',
                'about': '',
                'likes': '',
                'scraped_at': time.time()
            }
            
            # Facebook is heavily protected, this is basic extraction
            title_elem = soup.find('title')
            if title_elem:
                data['name'] = title_elem.get_text().strip()
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'platform': 'facebook'}
    
    def _scrape_youtube(self, username_or_url):
        """Scrape YouTube channel"""
        try:
            if not username_or_url.startswith('http'):
                url = f"https://youtube.com/@{username_or_url}"
            else:
                url = username_or_url
                
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'platform': 'youtube',
                'url': url,
                'name': '',
                'subscribers': '',
                'videos': '',
                'description': '',
                'scraped_at': time.time()
            }
            
            # Extract channel info
            name_elem = soup.find('meta', property='og:title')
            if name_elem:
                data['name'] = name_elem.get('content', '')
            
            desc_elem = soup.find('meta', property='og:description')
            if desc_elem:
                data['description'] = desc_elem.get('content', '')
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'platform': 'youtube'}
    
    def _scrape_tiktok(self, username_or_url):
        """Scrape TikTok profile"""
        try:
            if not username_or_url.startswith('http'):
                url = f"https://tiktok.com/@{username_or_url}"
            else:
                url = username_or_url
                
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'platform': 'tiktok',
                'url': url,
                'name': '',
                'username': '',
                'bio': '',
                'followers': '',
                'following': '',
                'likes': '',
                'scraped_at': time.time()
            }
            
            # TikTok uses JSON-LD structured data
            json_scripts = soup.find_all('script', type='application/ld+json')
            for script in json_scripts:
                try:
                    json_data = json.loads(script.string)
                    if '@type' in json_data and json_data['@type'] == 'Person':
                        data['name'] = json_data.get('name', '')
                        data['bio'] = json_data.get('description', '')
                        break
                except:
                    continue
            
            return data
            
        except Exception as e:
            return {'error': str(e), 'platform': 'tiktok'}
    
    def scrape_analytics_tools(self, tool_type, url_or_id):
        """Scrape analytics and SEO tools"""
        try:
            if tool_type == 'google_analytics':
                return self._scrape_google_analytics(url_or_id)
            elif tool_type == 'semrush':
                return self._scrape_semrush(url_or_id)
            elif tool_type == 'ahrefs':
                return self._scrape_ahrefs(url_or_id)
            else:
                return {'error': f'Unsupported analytics tool: {tool_type}'}
        except Exception as e:
            return {'error': str(e), 'tool': tool_type}
    
    def _scrape_google_analytics(self, property_id):
        """Scrape Google Analytics data (requires API access)"""
        # This would require Google Analytics API integration
        return {
            'tool': 'google_analytics',
            'property_id': property_id,
            'note': 'Requires Google Analytics API integration',
            'scraped_at': time.time()
        }
    
    def _scrape_semrush(self, domain):
        """Scrape SEMrush data"""
        try:
            url = f"https://www.semrush.com/analytics/overview/?q={domain}"
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'tool': 'semrush',
                'domain': domain,
                'url': url,
                'scraped_at': time.time()
            }
            
            # SEMrush requires authentication, this is basic structure
            return data
            
        except Exception as e:
            return {'error': str(e), 'tool': 'semrush'}
    
    def _scrape_ahrefs(self, domain):
        """Scrape Ahrefs data"""
        try:
            url = f"https://ahrefs.com/site-explorer/overview/v2/subdomains?target={domain}"
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            data = {
                'tool': 'ahrefs',
                'domain': domain,
                'url': url,
                'scraped_at': time.time()
            }
            
            # Ahrefs requires authentication, this is basic structure
            return data
            
        except Exception as e:
            return {'error': str(e), 'tool': 'ahrefs'}
    
    def process_document(self, file_path, file_type):
        """Process uploaded documents"""
        try:
            if file_type == 'pdf':
                return self._process_pdf(file_path)
            elif file_type == 'docx':
                return self._process_docx(file_path)
            elif file_type in ['jpg', 'jpeg', 'png', 'gif']:
                return self._process_image(file_path)
            elif file_type == 'txt':
                return self._process_text(file_path)
            else:
                return {'error': f'Unsupported file type: {file_type}'}
        except Exception as e:
            return {'error': str(e), 'file_type': file_type}
    
    def _process_pdf(self, file_path):
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    'file_type': 'pdf',
                    'text': text,
                    'pages': len(pdf_reader.pages),
                    'processed_at': time.time()
                }
        except Exception as e:
            return {'error': str(e), 'file_type': 'pdf'}
    
    def _process_docx(self, file_path):
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'file_type': 'docx',
                'text': text,
                'paragraphs': len(doc.paragraphs),
                'processed_at': time.time()
            }
        except Exception as e:
            return {'error': str(e), 'file_type': 'docx'}
    
    def _process_image(self, file_path):
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                'file_type': 'image',
                'text': text,
                'dimensions': image.size,
                'processed_at': time.time()
            }
        except Exception as e:
            return {'error': str(e), 'file_type': 'image'}
    
    def _process_text(self, file_path):
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'file_type': 'text',
                'text': text,
                'characters': len(text),
                'processed_at': time.time()
            }
        except Exception as e:
            return {'error': str(e), 'file_type': 'text'}
    
    def scrape_all_user_data(self, user_data):
        """Scrape all available data for a user"""
        scraped_data = {
            'website': {},
            'social_media': {},
            'analytics_tools': {},
            'documents': {},
            'scraped_at': time.time()
        }
        
        # Scrape website if provided
        if user_data.get('websiteUrl'):
            scraped_data['website'] = self.scrape_website(user_data['websiteUrl'])
        
        # Scrape social media links
        social_platforms = ['linkedinPage', 'twitterHandle', 'instagramAccount', 'facebookPage', 'tiktokYoutube']
        for platform in social_platforms:
            if user_data.get(platform):
                platform_name = platform.replace('Page', '').replace('Handle', '').replace('Account', '').lower()
                if platform_name == 'tiktokyoutube':
                    # Handle both TikTok and YouTube
                    if 'tiktok' in user_data[platform].lower():
                        scraped_data['social_media']['tiktok'] = self.scrape_social_media('tiktok', user_data[platform])
                    elif 'youtube' in user_data[platform].lower():
                        scraped_data['social_media']['youtube'] = self.scrape_social_media('youtube', user_data[platform])
                else:
                    scraped_data['social_media'][platform_name] = self.scrape_social_media(platform_name, user_data[platform])
        
        # Scrape analytics tools
        analytics_tools = ['googleAnalytics', 'seoTools', 'ecommercePlatforms', 'adPlatforms']
        for tool in analytics_tools:
            if user_data.get(tool):
                tool_name = tool.lower()
                scraped_data['analytics_tools'][tool_name] = self.scrape_analytics_tools(tool_name, user_data[tool])
        
        return scraped_data
